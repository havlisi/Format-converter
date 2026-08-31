import builtins
import os
import sys
import pytest
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from core import docx_io
from core.docx_io import read_docx, write_docx


def _make_text_pdf(path):
    doc = SimpleDocTemplate(path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = [
        Paragraph("Kontoauszug Nr. 7", styles["Title"]),
        Paragraph("Kontoinhaber: Musterfirma GmbH", styles["Normal"]),
        Spacer(1, 12),
        Table([["Datum", "Betrag"], ["01.03.2026", "1.234,56"], ["02.03.2026", "-99,00"]]),
    ]
    doc.build(story)


def test_pdf_to_docx_preserves_text_and_table(tmp_path):
    pdf = str(tmp_path / "src.pdf")
    out = str(tmp_path / "src.docx")
    _make_text_pdf(pdf)

    docx_io.pdf_to_docx(pdf, out)

    assert os.path.exists(out)
    doc = Document(out)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    all_text += "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "Kontoinhaber: Musterfirma GmbH" in all_text
    assert "1.234,56" in all_text
    assert "02.03.2026" in all_text


def test_pdf_to_docx_scanned_falls_back_to_block_pipeline(tmp_path, monkeypatch):
    pdf = str(tmp_path / "scan.pdf")
    out = str(tmp_path / "scan.docx")
    # content doesn't matter — we force the "no text layer" branch
    open(pdf, "wb").write(b"%PDF-1.4\n%%EOF\n")

    monkeypatch.setattr(docx_io, "_pdf_has_text_layer", lambda p: False)
    sentinel = [("text", "OCR line one"), ("table", [["a", "b"]])]
    monkeypatch.setattr("core.pdf_io.read_pdf", lambda p: sentinel)

    called = {}
    real_write = docx_io.write_docx

    def spy_write(blocks, path):
        called["blocks"] = blocks
        real_write(blocks, path)

    monkeypatch.setattr(docx_io, "write_docx", spy_write)

    docx_io.pdf_to_docx(pdf, out)

    assert called["blocks"] == sentinel
    assert os.path.exists(out)


def test_pdf_to_docx_import_systemexit_becomes_runtimeerror(tmp_path, monkeypatch):
    # F2: `from pdf2docx import Converter` can raise SystemExit at import time
    # (a PyMuPDF version check inside pdf2docx). SystemExit is a BaseException,
    # so batch.run_batch's `except Exception` would NOT catch it — one bad
    # environment would kill the whole batch. The import must be guarded so it
    # surfaces as a normal, catchable RuntimeError instead.
    #
    # Approach: patch builtins.__import__ so importing pdf2docx raises
    # SystemExit, and force the text-layer branch so the import is reached.
    pdf = str(tmp_path / "src.pdf")
    out = str(tmp_path / "src.docx")
    open(pdf, "wb").write(b"%PDF-1.4\n%%EOF\n")

    monkeypatch.setattr(docx_io, "_pdf_has_text_layer", lambda p: True)

    real_import = builtins.__import__

    def fake_import(name, *args, **kwargs):
        if name == "pdf2docx" or name.startswith("pdf2docx."):
            raise SystemExit("boom")
        return real_import(name, *args, **kwargs)

    monkeypatch.delitem(sys.modules, "pdf2docx", raising=False)
    monkeypatch.setattr(builtins, "__import__", fake_import)

    with pytest.raises(RuntimeError):
        docx_io.pdf_to_docx(pdf, out)


def test_pdf_has_text_layer_false_for_sparse_stamp_only_pdf(tmp_path):
    # F4: a scanned page whose only text layer is a 2-word Bates stamp / fax
    # banner must still route to OCR — a mean of ~2 words/page is below the
    # density threshold, so _pdf_has_text_layer returns False.
    pdf = str(tmp_path / "stamp.pdf")
    doc = SimpleDocTemplate(pdf, pagesize=letter)
    doc.build([Paragraph("EXHIBIT 12", getSampleStyleSheet()["Normal"])])

    assert docx_io._pdf_has_text_layer(pdf) is False


def test_pdf_has_text_layer_true_for_dense_text_pdf(tmp_path):
    # F4 companion: a real text layer (well over 5 words/page) still takes the
    # pdf2docx branch.
    pdf = str(tmp_path / "dense.pdf")
    _make_text_pdf(pdf)

    assert docx_io._pdf_has_text_layer(pdf) is True


def test_write_then_read_text_block(tmp_path):
    p = tmp_path / "out.docx"
    blocks = [("text", "Hello world")]

    write_docx(blocks, str(p))
    result = read_docx(str(p))

    assert result == [("text", "Hello world")]


def test_write_then_read_table_block(tmp_path):
    p = tmp_path / "out2.docx"
    blocks = [("table", [["Name", "Age"], ["Bob", "30"]])]

    write_docx(blocks, str(p))
    result = read_docx(str(p))

    assert result == [("table", [["Name", "Age"], ["Bob", "30"]])]


def test_order_preserved_text_then_table(tmp_path):
    p = tmp_path / "out3.docx"
    blocks = [("text", "Intro"), ("table", [["x", "y"]])]

    write_docx(blocks, str(p))
    result = read_docx(str(p))

    assert result == [("text", "Intro"), ("table", [["x", "y"]])]
