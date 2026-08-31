from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from core.types import Block
from typing import List


def read_docx(path: str) -> List[Block]:
    doc = Document(path)
    blocks: List[Block] = []
    for element in doc.element.body.iterchildren():
        if element.tag.endswith("}p"):
            para = Paragraph(element, doc)
            if para.text.strip():
                blocks.append(("text", para.text))
        elif element.tag.endswith("}tbl"):
            table = Table(element, doc)
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            blocks.append(("table", rows))
    return blocks


def write_docx(blocks: List[Block], path: str) -> None:
    doc = Document()
    for kind, content in blocks:
        if kind == "text":
            doc.add_paragraph(content)
        else:
            if not content:
                continue
            n_rows = len(content)
            n_cols = max(len(r) for r in content)
            table = doc.add_table(rows=n_rows, cols=n_cols)
            table.style = "Table Grid"
            for r_idx, row in enumerate(content):
                for c_idx in range(n_cols):
                    val = row[c_idx] if c_idx < len(row) else ""
                    table.cell(r_idx, c_idx).text = "" if val is None else str(val)
    doc.save(path)


# A real text layer averages far more than this per page; a scanned statement
# that only carries a one-word Bates stamp or a fax banner on an otherwise
# image-only page falls well below it and must still go through OCR — as must a
# mostly-scanned PDF with a stray text page or two among dozens of image pages.
_MIN_MEAN_WORDS_PER_PAGE = 5


def _pdf_has_text_layer(pdf_path: str) -> bool:
    """True only if the PDF carries a *meaningful* text layer — mean extractable
    words per page above _MIN_MEAN_WORDS_PER_PAGE. A scanned/image-only PDF (or one
    whose only text is a stray stamp/banner) stays below that and is routed through
    OCR instead of pdf2docx, so its scanned pages aren't silently lost."""
    import pdfplumber

    total_words = 0
    page_count = 0
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_count += 1
            total_words += len(page.extract_words())
    return page_count > 0 and total_words > _MIN_MEAN_WORDS_PER_PAGE * page_count


def pdf_to_docx(pdf_path: str, docx_path: str) -> None:
    """Convert a PDF to DOCX preserving page layout via pdf2docx. A scanned PDF
    with no text layer has no layout to preserve — route it through the existing
    OCR block pipeline so the DOCX at least carries searchable text."""
    if not _pdf_has_text_layer(pdf_path):
        from core import pdf_io
        write_docx(pdf_io.read_pdf(pdf_path), docx_path)
        return

    import logging

    # pdf2docx calls logging.basicConfig(level=INFO) at import time, attaching a
    # handler to the ROOT logger and forcing root to INFO process-wide — snapshot
    # root and restore it afterward so that side effect doesn't leak. The import
    # itself can also raise SystemExit (a PyMuPDF version check inside pdf2docx);
    # SystemExit is a BaseException, so batch.run_batch's `except Exception`
    # wouldn't catch it — one bad environment would kill the whole batch. Convert
    # it (and a plain ImportError) into a normal, catchable RuntimeError.
    _root = logging.getLogger()
    _prev_level = _root.level
    _prev_handlers = list(_root.handlers)
    try:
        from pdf2docx import Converter
    except (ImportError, SystemExit) as e:
        raise RuntimeError(f"pdf2docx unavailable: {e}") from e
    _root.setLevel(_prev_level)
    for _h in list(_root.handlers):
        if _h not in _prev_handlers:
            _root.removeHandler(_h)

    cv = Converter(pdf_path)
    try:
        cv.convert(docx_path)
    finally:
        cv.close()
